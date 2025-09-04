"""
Updated: supports large tables using LongTable + docx export.
Processor module.
Expose: generate_reports_from_csv(input_csv: str, out_dir: str) -> dict
Produces: out_dir/analysis_output.csv, out_dir/report.pdf, out_dir/report.docx (optional)
"""

import os,re,sys,csv,logging
from datetime import datetime
from pathlib import Path
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from wordcloud import WordCloud, STOPWORDS
from transformers import pipeline
from sklearn.feature_extraction.text import CountVectorizer
from sklearn.decomposition import LatentDirichletAllocation
  
# reportlab platypus
from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, PageBreak,
                                TableStyle, Image, LongTable)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_LEFT

# try import python-docx (optional)
DOCX_AVAILABLE = True
try:
    from docx import Document
    from docx.shared import Inches
except Exception:
    DOCX_AVAILABLE = False

logger = logging.getLogger("processor")
logger.setLevel(logging.INFO)

# ---------------- CONFIG ----------------
CSV_ENCODING = "utf-8"
MAX_ROWS = None          # None => all rows
TOPIC_COUNT = 3

# Table teaser length to avoid massive single-cell height in PDF tables
TEASER_CHAR_LIMIT = 900

# ---------------- UTIL ----------------
RELATIVE_TIME_RE = re.compile(
    r'(?:(\d+)\s*(second|sec|s|minute|min|m|hour|hr|h|day|d|week|w|month|mo|year|yr|y)s?\s*ago)|\b(yesterday|today|just now|now)\b',
    flags=re.IGNORECASE
)

try:
    import torch
    device = 0 if torch.cuda.is_available() else -1
except Exception:
    device = -1

try:
    sentiment_model = pipeline("sentiment-analysis",
                            model="distilbert-base-uncased-finetuned-sst-2-english",
                            device=device)
except Exception as e:
    print("Failed to load requested model:", e)
    try:
        sentiment_model = pipeline("sentiment-analysis", device=device)
    except Exception as ex:
        print("Final sentiment pipeline fallback failed:", ex); sys.exit(1)
            

def parse_relative_time(s: str, ref: pd.Timestamp):
    if not isinstance(s, str) or s.strip() == "":
        return pd.NaT
    s = s.strip().lower()
    if s in ("just now", "now"):
        return ref
    if s == "today":
        return pd.Timestamp(ref.date())
    if s == "yesterday":
        return ref - pd.Timedelta(days=1)
    s = re.sub(r'\b(an|a)\b', '1', s)
    m = re.search(r'(\d+)\s*(second|sec|s|minute|min|m|hour|hr|h|day|d|week|w|month|mo|year|yr|y)s?\s*ago', s)
    if not m:
        return pd.NaT
    qty = int(m.group(1)); unit = m.group(2).lower()
    if unit in ("second","sec","s"): return ref - pd.Timedelta(seconds=qty)
    if unit in ("minute","min","m"): return ref - pd.Timedelta(minutes=qty)
    if unit in ("hour","hr","h"): return ref - pd.Timedelta(hours=qty)
    if unit in ("day","d"): return ref - pd.Timedelta(days=qty)
    if unit in ("week","w"): return ref - pd.Timedelta(weeks=qty)
    if unit in ("month","mo"): return ref - pd.Timedelta(days=qty * 30)
    if unit in ("year","yr","y"): return ref - pd.Timedelta(days=qty * 365)
    return pd.NaT

def clean_text(text: str) -> str:
    if not isinstance(text, str): return ""
    text = re.sub(r"http\S+", "", text)
    text = re.sub(r"@\w+", "", text)
    text = re.sub(r"#\w+", "", text)
    text = re.sub(r"[^A-Za-z\s]", " ", text)
    text = re.sub(r"\s+", " ", text)
    return text.lower().strip()

def chunked(iterable, size):
    for i in range(0, len(iterable), size):
        yield iterable[i:i+size]


def teaser(s, n=TEASER_CHAR_LIMIT):
    if not isinstance(s, str): return ""
    s = s.strip()
    return (s if len(s) <= n else s[:n-1].rsplit(" ",1)[0] + " ...")

def parse_score(x):
    if pd.isna(x): return np.nan
    s = str(x)
    m = re.search(r"(-?\d+)", s.replace(",", ""))
    if m: return int(m.group(1))
    nums = re.findall(r"\d+", s)
    return int(nums[0]) if nums else np.nan

def parse_time_value(v,ref_ts):
    if isinstance(v, (pd.Timestamp, datetime)): return pd.to_datetime(v)
    if pd.isna(v): return pd.NaT
    s = str(v).strip()
    try:
        parsed = pd.to_datetime(s, errors='coerce', utc=None)
        if pd.notna(parsed): return parsed
    except Exception: pass
    rt = parse_relative_time(s, ref_ts)
    if pd.notna(rt): return pd.to_datetime(rt)
    return pd.NaT

def compile_list(lst): return [re.compile(pat, flags=re.IGNORECASE) for pat in lst]


# ---------------- India-specific nature detection ----------------
PRO_INDIA = [r"\bjai hind\b", r"\bvande mataram\b", r"\bpro india\b", r"\bpro-india\b", r"\bsupport (?:india|modi|bjp)\b", r"\bproud of india\b", r"\bindia is great\b"]
ANTI_INDIA = [r"\banti[- ]?india\b", r"\banti national\b", r"\btraitor\b", r"\banti-india\b", r"\bkill india\b", r"\bboycott india\b"]
CRITICAL_GOVT = [r"\bmodi sucks\b", r"\bcorrupt government\b", r"\bgovernment (?:is )?failing\b", r"\b(criticis|criticize|criticising) (?:government|modi|bjp)\b", r"\bpolicy (?:failure|fail)\b", r"\banti-corruption\b", r"\bmisgovern(ance|ing)\b", r"\bgovernment (?:policy|policies)"]
SUPPORT_OPPOSITION = [r"\bsupport (?:congress|aam aadmi|aap|opposition)\b", r"\bvot(e|ing) for .*opposition\b"]
SEPARATIST = [r"\bazadi\b", r"\bseparatist\b", r"\bsecede\b", r"\bindependence for\b"]
COMMUNAL = [r"\bcommunal\b", r"\breligious (?:tension|hatred)\b", r"\breligious\b", r"\bminority\b"]
CALL_TO_ACTION = [r"\bprotest\b", r"\bboycott\b", r"\bjoin (?:the )?protest\b", r"\bstrike\b", r"\brally\b", r"\baction\b"]
CONSPIRACY = [r"\bforeign funded\b", r"\bdeep state\b", r"\bconspiracy\b", r"\bwestern plot\b", r"\bcia\b", r"\bsecret agenda\b"]

PRO_INDIA_RE = compile_list(PRO_INDIA); ANTI_INDIA_RE = compile_list(ANTI_INDIA)
CRITICAL_GOVT_RE = compile_list(CRITICAL_GOVT); SUPPORT_OPPOSITION_RE = compile_list(SUPPORT_OPPOSITION)
SEPARATIST_RE = compile_list(SEPARATIST); COMMUNAL_RE = compile_list(COMMUNAL)
CALL_TO_ACTION_RE = compile_list(CALL_TO_ACTION); CONSPIRACY_RE = compile_list(CONSPIRACY)


def text_matches_any(text, patterns):
    for pat in patterns:
        if pat.search(text or ""): return True
    return False

def determine_nature(text, sentiment_label):
    t = (text or "").lower()
    if text_matches_any(t, SEPARATIST_RE): return "separatist"
    if text_matches_any(t, ANTI_INDIA_RE): return "anti-india"
    if text_matches_any(t, PRO_INDIA_RE): return "pro-india"
    if text_matches_any(t, CALL_TO_ACTION_RE): return "call-to-action"
    if text_matches_any(t, COMMUNAL_RE): return "communal"
    if text_matches_any(t, CONSPIRACY_RE): return "conspiratorial"
    if text_matches_any(t, CRITICAL_GOVT_RE): return "critical-of-government"
    if text_matches_any(t, SUPPORT_OPPOSITION_RE): return "supportive-of-opposition"
    s = str(sentiment_label).upper()
    if "POS" in s: return "supportive"
    if "NEG" in s: return "critical"
    return "neutral"

# ---------------- DANGEROUS FLAG ----------------
danger_keywords = ["kill","attack","bomb","violence","terror","terrorist","militant","insurgency","boycott","protest","call to action"]
pattern = re.compile(r'\b(?:' + '|'.join(map(re.escape, danger_keywords)) + r')\b', flags=re.IGNORECASE)

def is_dangerous(text, sentiment):
    if pattern.search(text or ""): return True
    return (str(sentiment).upper() == "NEGATIVE" and text.strip() != "")

def generate_reports_from_csv(input_csv:str, out_dir:str) -> dict:
    """
    Runs full analysis pipeline. Returns dict: {'pdf':..., 'csv':..., 'docx':...}
    """
    logger.info("Running processing pipeline on %s",input_csv)
    out_dir= Path(out_dir)
    out_dir.mkdir(parents=True,exist_ok=True)

    # ---------------- READ CSV ----------------
    if not os.path.exists(input_csv):
        print("CSV file not found:", input_csv); sys.exit(1)

    print("Loading CSV:", input_csv)
    try:
        df_raw = pd.read_csv(input_csv, encoding=CSV_ENCODING, low_memory=False)
    except Exception as e:
        print("Error reading CSV:", e); sys.exit(1)

    if MAX_ROWS:
        df_raw = df_raw.head(MAX_ROWS)

    title_col = "Title"
    reference_col = "Reference"
    subreddit_col = "Subreddit"
    score_col = "Score"
    comment_col = "Comments"
    time_col = "Time"
    author_col = "Author"
    desc_col = "Description"
    url_col = "Url"

    if not any(c in df_raw.columns for c in [title_col, comment_col, desc_col]):
        print("No text column detected. CSV columns:", list(df_raw.columns)); sys.exit(1)

# if title is None(not provided) entire column is filled with "" strings
# if title is provided but for some it is NaN after astype(str) they become "nan" not empty string
    # normalized df
    df = pd.DataFrame()
    df["orig_index"] = df_raw.index.astype(str)
    df["title"] = df_raw[title_col].fillna("").astype(str) if title_col else ""
    df["reference"] = df_raw[reference_col].astype(str) if reference_col else ""
    df["subreddit"] = df_raw[subreddit_col] if subreddit_col else "N/A"
    df["raw_score"] = df_raw[score_col] if score_col else np.nan
    df["comment"] = df_raw[comment_col].fillna("").astype(str) if comment_col else ""
    df["time_raw"] = df_raw[time_col] if time_col else ""
    df["username"] = df_raw[author_col] if author_col else "N/A"
    df["description"] = df_raw[desc_col].fillna("").astype(str) if desc_col else ""
    df["url"] = df_raw[url_col] if url_col else ""
  
    df["text_for_analysis"] = (df["title"] + " " + df["comment"] + " " + df["description"]).str.strip()
    df.loc[df["text_for_analysis"].str.strip() == "", "text_for_analysis"] = df.loc[df["text_for_analysis"].str.strip() == "", :].apply(
        lambda r: " ".join([str(v) for v in r.values if isinstance(v, str) and v.strip() != ""]), axis=1
    )
    df["clean_text"] = df["text_for_analysis"].apply(clean_text)
    df["score"] = df["raw_score"].apply(parse_score)

    # parse times
    try:
        ref_ts = pd.to_datetime(os.path.getmtime(input_csv), unit='s')
    except Exception:
        ref_ts = pd.Timestamp.now()
    
    df["created_at"] = df["time_raw"].apply(lambda x: parse_time_value(x,ref_ts))

    # ---------------- SENTIMENT ----------------
    print("Loading sentiment model...")

    texts = df["clean_text"].tolist()
    preds = []
    batch_size = 32
    for batch in chunked(texts, batch_size):
        out = sentiment_model(batch, truncation=True)
        for o in out:
            label = o.get("label", "NEUTRAL")
            score = float(o.get("score", 0.0))
            preds.append((label, score))

    df["sentiment"] = [p[0] for p in preds]
    df["sentiment_score"] = [p[1] for p in preds]

    df["nature"] = df.apply(lambda r: determine_nature(r["clean_text"], r["sentiment"]), axis=1)

    # ---------------- TOPIC MODELING ----------------
    print("Performing topic modeling...")

    vectorizer = CountVectorizer(stop_words="english", min_df=2)
    try:
        X = vectorizer.fit_transform(df["clean_text"])
    except Exception as e:
        print("Topic vectorization failed:", e); X = None

    if X is None or X.shape[0] < 3 or len(vectorizer.get_feature_names_out()) < 5:
        df["topic"] = np.nan
        topic_counts = pd.Series(dtype=int)
    else:
        n_topics = min(TOPIC_COUNT, X.shape[0])
        lda = LatentDirichletAllocation(n_components=n_topics, random_state=42)
        lda.fit(X)
        doc_topic = lda.transform(X)
        df["topic"] = doc_topic.argmax(axis=1)
        topic_counts = df["topic"].value_counts().sort_index()

    df["dangerous"] = df.apply(lambda r: is_dangerous(r["clean_text"], r["sentiment"]), axis=1)
    dangerous_tweets = df[df["dangerous"]].copy()
    print(f"Flagged {len(dangerous_tweets)} potentially dangerous posts.")

    # ---------------- VISUALS ----------------
    try:
        # sentiment plot
        sent_counts = df["sentiment"].value_counts()
        plt.figure(figsize=(6,4))
        sent_counts.plot(kind="bar")
        plt.title("Sentiment Distribution")
        plt.tight_layout()
        plt.savefig(out_dir / "sentiment.png", dpi=150)
        plt.close()
        # topic plot
        if "topic" in df and df["topic"].notna().any():
            topic_counts = df["topic"].value_counts().sort_index()
            plt.figure(figsize=(6,4))
            topic_counts.plot(kind="bar")
            plt.title("Topic Distribution")
            plt.tight_layout()
            plt.savefig(out_dir / "topics.png", dpi=150)
            plt.close()
        # danger wordcloud
        dangerous_df = df[df["dangerous"]]
        if not dangerous_df.empty:
            wc_text = " ".join(dangerous_df["clean_text"].tolist())
            wc = WordCloud(width=1000, height=400, background_color="white", stopwords=set(STOPWORDS)).generate(wc_text)
            plt.figure(figsize=(12,5))
            plt.imshow(wc, interpolation="bilinear")
            plt.axis("off")
            plt.tight_layout()
            plt.savefig(out_dir / "danger_wc.png", dpi=150)
            plt.close()
    except Exception as e:
        logger.warning("Visuals generation failed: %s", e)


    # ---------------- BUILD PDF ----------------
    print("Building PDF report (LongTable for large tables)...")
    pdf_out= out_dir/"report.pdf"
    styles = getSampleStyleSheet()
    styleN = styles["Normal"]
    styleH = styles["Heading2"]
    title_style = styles["Title"]
    tweet_paragraph_style = ParagraphStyle("TweetStyle", parent=styles["BodyText"], fontSize=9, leading=11, spaceAfter=6, alignment=TA_LEFT)

    doc = SimpleDocTemplate(pdf_out, pagesize=A4, rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=36)
    elements = []
    elements.append(Paragraph("Reddit Posts Report (CSV Source) — India-specific Nature", title_style))
    elements.append(Spacer(1, 8))
    elements.append(Paragraph(f"Total Posts Processed: {len(df)}", styleN))
    elements.append(Spacer(1, 8))

    # Sentiment summary
    elements.append(Paragraph("Sentiment Analysis Summary", styleH))
    total = len(df)
    for label, count in sent_counts.items():
        pct = count / total * 100 if total > 0 else 0
        elements.append(Paragraph(f"{label}: {count} posts ({pct:.1f}%)", styleN))
    elements.append(Spacer(1, 6))
    if os.path.exists("sentiment.png"):
        elements.append(Image("sentiment.png", width=5.5*inch, height=3*inch))
    elements.append(Spacer(1, 12))

    # Topic & Nature summary
    if not topic_counts.empty:
        elements.append(Paragraph("Topic Modeling Summary", styleH))
        for idx, val in topic_counts.items():
            elements.append(Paragraph(f"Topic {int(idx)}: {int(val)} posts", styleN))
        elements.append(Spacer(1, 6))
        if os.path.exists("topics.png"): elements.append(Image("topics.png", width=5.5*inch, height=3*inch))
        elements.append(Spacer(1, 12))

    elements.append(Paragraph("Nature (India-specific) Summary", styleH))
    nature_counts = df["nature"].value_counts()
    for label, count in nature_counts.items():
        pct = count / total * 100 if total > 0 else 0
        elements.append(Paragraph(f"{label}: {count} posts ({pct:.1f}%)", styleN))
    elements.append(Spacer(1, 12))

    # Dangerous posts table (LongTable)
    elements.append(Paragraph("Flagged Potentially Dangerous Posts", styleH))
    elements.append(Spacer(1, 6))
    if dangerous_tweets.empty:
        elements.append(Paragraph("No dangerous posts detected.", styleN))
    else:
        # prepare LongTable data (header + rows)
        header = ["Post (teaser)", "Subreddit", "Author", "Sentiment", "Nature", "Topic", "Date"]
        lt_data = [header]
        for _, row in dangerous_tweets.iterrows():
            date_str = row["created_at"].strftime("%Y-%m-%d %H:%M") if pd.notna(row["created_at"]) else "N/A"
            lt_data.append([
                Paragraph(teaser(row["text_for_analysis"], TEASER_CHAR_LIMIT), tweet_paragraph_style),
                row["subreddit"] if pd.notna(row["subreddit"]) else "N/A",
                row["username"] if pd.notna(row["username"]) else "N/A",
                row["sentiment"],
                row["nature"],
                str(int(row["topic"])) if not pd.isna(row["topic"]) else "N/A",
                date_str
            ])
        col_widths = [3.0*inch, 0.7*inch, 0.8*inch, 0.6*inch, 0.8*inch, 0.5*inch, 1.0*inch]
        lt = LongTable(lt_data, colWidths=col_widths, repeatRows=1)
        # style: small font, grid, header background
        lt_style = TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#4F81BD")),
            ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
            ('ALIGN', (1,0), (-1,-1), 'CENTER'),
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ('GRID', (0,0), (-1,-1), 0.25, colors.grey),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('FONTSIZE', (0,0), (-1,-1), 8),
            ('LEFTPADDING', (0,0), (-1,-1), 4),
            ('RIGHTPADDING', (0,0), (-1,-1), 4),
        ])
        lt.setStyle(lt_style)
        elements.append(lt)
        elements.append(Spacer(1, 12))
        if os.path.exists("danger_wc.png"):
            elements.append(Paragraph("Word Cloud of Flagged Posts", styleH)); elements.append(Image("danger_wc.png", width=5.5*inch, height=2.6*inch))

    elements.append(PageBreak())

    # All collected posts (LongTable) - include full dataset but use teaser to avoid huge cells
    elements.append(Paragraph("All Collected Posts", styles['Heading2']))
    all_header = ["Date", "Subreddit", "Author", "Score", "Nature", "Post (teaser)"]
    all_lt_data = [all_header]
    for idx, row in df.iterrows():
        date_str = row["created_at"].strftime("%Y-%m-%d %H:%M") if pd.notna(row["created_at"]) else "N/A"
        all_lt_data.append([
            date_str,
            row["subreddit"] if pd.notna(row["subreddit"]) else "N/A",
            row["username"] if pd.notna(row["username"]) else "N/A",
            str(row["score"]) if not pd.isna(row["score"]) else "N/A",
            row["nature"],
            Paragraph(teaser(row["text_for_analysis"], TEASER_CHAR_LIMIT), tweet_paragraph_style)
        ])

    all_col_widths = [1.0*inch, 1.0*inch, 1.0*inch, 0.7*inch, 0.9*inch, 2.8*inch]
    all_lt = LongTable(all_lt_data, colWidths=all_col_widths, repeatRows=1)
    all_lt.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#4F81BD")),
        ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
        ('GRID', (0,0), (-1,-1), 0.25, colors.grey),
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ('FONTSIZE', (0,0), (-1,-1), 8),
        ('LEFTPADDING', (0,0), (-1,-1), 4),
        ('RIGHTPADDING', (0,0), (-1,-1), 4),
    ]))
    elements.append(all_lt)

    # finish PDF
    doc = SimpleDocTemplate(str(pdf_out))
    doc.build(elements)
    print("✅ PDF saved as:", pdf_out)

    # ---------------- SAVE CSV (full enriched) ----------------
    csv_out = out_dir/"analysis_output.csv"
    df_out = df.copy()
    df_out["created_at_str"] = df_out["created_at"].apply(lambda x: x.strftime("%Y-%m-%d %H:%M:%S") if pd.notna(x) else "")
    df_out.to_csv(csv_out, index=False, encoding="utf-8")
    print("✅ Enriched CSV saved as:", csv_out)


    # ---------------- DOCX EXPORT (optional) ----------------
    if not DOCX_AVAILABLE:
        print("python-docx not installed — skipping DOCX export. Install via: pip install python-docx")
    else:
        try:
            print("Building DOCX report...")
            DOCX_OUTPUT= out_dir/"report.docx"
            docx = Document()
            docx.add_heading("Reddit Posts Report (India-specific Nature)", level=1)
            docx.add_paragraph(f"Total Posts Processed: {len(df)}")
            docx.add_heading("Sentiment Analysis Summary", level=2)
            for label, count in sent_counts.items():
                pct = count / total * 100 if total > 0 else 0
                docx.add_paragraph(f"{label}: {count} posts ({pct:.1f}%)")

            docx.add_heading("Nature Summary", level=2)
            for label, count in nature_counts.items():
                pct = count / total * 100 if total > 0 else 0
                docx.add_paragraph(f"{label}: {count} posts ({pct:.1f}%)")

            # add small sample table (first 200 rows or less)
            sample_n = min(200, len(df))
            docx.add_heading(f"Sample of First {sample_n} Posts", level=2)
            table = docx.add_table(rows=1, cols=6)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Date"
            hdr_cells[1].text = "Subreddit"
            hdr_cells[2].text = "Author"
            hdr_cells[3].text = "Score"
            hdr_cells[4].text = "Nature"
            hdr_cells[5].text = "Post (teaser)"
            for idx, row in df.head(sample_n).iterrows():
                row_cells = table.add_row().cells
                date_str = row["created_at"].strftime("%Y-%m-%d %H:%M") if pd.notna(row["created_at"]) else "N/A"
                row_cells[0].text = date_str
                row_cells[1].text = str(row["subreddit"]) if pd.notna(row["subreddit"]) else "N/A"
                row_cells[2].text = str(row["username"]) if pd.notna(row["username"]) else "N/A"
                row_cells[3].text = str(row["score"]) if not pd.isna(row["score"]) else "N/A"
                row_cells[4].text = str(row["nature"])
                row_cells[5].text = teaser(row["text_for_analysis"], 300)

            docx.save(DOCX_OUTPUT)
            print("✅ DOCX saved as:", DOCX_OUTPUT)
        except Exception as e:
            logger.exception("DOCX creation failed: %s", e)
            if DOCX_OUTPUT.exists():
                try:
                    DOCX_OUTPUT.unlink(missing_ok=True)
                except Exception:
                    pass
    logger.info("Processor: finished, files at %s", out_dir)
    return {"pdf": str(pdf_out), "csv": str(csv_out), "docx": str(DOCX_OUTPUT) if DOCX_OUTPUT.exists() else ""}
    
